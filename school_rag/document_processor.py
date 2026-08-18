"""
文档处理模块 - 支持多种格式的文档解析和文本提取

支持格式：PDF, DOCX, XLSX, TXT, MD, 图片（OCR）
"""

import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import re
from loguru import logger


class DocumentProcessor:
    """文档处理器 - 解析和切分文档"""

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        ".pdf", ".docx", ".doc", ".xlsx", ".xls",
        ".txt", ".md", ".markdown", ".jpg", ".jpeg", ".png", ".bmp"
    }

    # 文档切分配置
    DEFAULT_CHUNK_SIZE = 500  # 字符数
    DEFAULT_CHUNK_OVERLAP = 100  # 重叠字符数
    DEFAULT_MIN_CHUNK_SIZE = 50  # 最小块大小

    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
        min_chunk_size: int = None,
    ):
        """
        初始化文档处理器

        Args:
            chunk_size: 切分块大小
            chunk_overlap: 块之间重叠大小
            min_chunk_size: 最小块大小
        """
        self.chunk_size = chunk_size or self.DEFAULT_CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or self.DEFAULT_CHUNK_OVERLAP
        self.min_chunk_size = min_chunk_size or self.DEFAULT_MIN_CHUNK_SIZE

        # 初始化 OCR 引擎（可选）
        self.ocr_available = self._init_ocr()

    def _init_ocr(self) -> bool:
        """初始化 OCR 引擎"""
        try:
            import paddleocr
            self.paddle_ocr = paddleocr.PaddleOCR(
                use_angle_cls=True,
                lang="ch",
                show_log=False
            )
            logger.info("PaddleOCR initialized successfully")
            return True
        except ImportError:
            logger.warning("PaddleOCR not available, image processing will be limited")
            return False
        except Exception as e:
            logger.warning(f"PaddleOCR initialization failed: {e}")
            return False

    def is_supported_format(self, file_path: str) -> bool:
        """
        检查文件格式是否支持

        Args:
            file_path: 文件路径

        Returns:
            是否支持
        """
        return Path(file_path).suffix.lower() in self.SUPPORTED_FORMATS

    def extract_text_from_file(self, file_path: str) -> str:
        """
        从文件中提取文本

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        try:
            if suffix == ".pdf":
                return self._extract_from_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                return self._extract_from_docx(file_path)
            elif suffix in [".xlsx", ".xls"]:
                return self._extract_from_excel(file_path)
            elif suffix in [".txt", ".md", ".markdown"]:
                return self._extract_from_text(file_path)
            elif suffix in [".jpg", ".jpeg", ".png", ".bmp"]:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_from_pdf(self, file_path: Path) -> str:
        """从 PDF 提取文本"""
        try:
            # 优先使用 pdfplumber（更好的中文支持）
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n\n".join(text_parts)

        except ImportError:
            # 回退到 PyPDF2
            try:
                import PyPDF2

                text_parts = []
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text_parts.append(page.extract_text())

                return "\n\n".join(text_parts)

            except ImportError:
                logger.error("Neither pdfplumber nor PyPDF2 is available")
                raise ImportError("Please install pdfplumber or PyPDF2 for PDF processing")

    def _extract_from_docx(self, file_path: Path) -> str:
        """从 DOCX 提取文本"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("python-docx is not available")
            raise ImportError("Please install python-docx for DOCX processing")

    def _extract_from_excel(self, file_path: Path) -> str:
        """从 Excel 提取文本"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []

                for row in sheet.iter_rows(values_only=True):
                    # 跳过空行
                    if any(cell for cell in row if cell is not None):
                        row_text = " | ".join([
                            str(cell) if cell is not None else "" for cell in row
                        ])
                        sheet_text.append(row_text)

                if sheet_text:
                    text_parts.append(f"=== 表格: {sheet_name} ===")
                    text_parts.extend(sheet_text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("openpyxl is not available")
            raise ImportError("Please install openpyxl for Excel processing")

    def _extract_from_text(self, file_path: Path) -> str:
        """从文本文件提取内容"""
        try:
            # 尝试多种编码
            encodings = ["utf-8", "gbk", "gb2312", "utf-16"]

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            # 如果所有编码都失败，使用错误忽略模式
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise

    def _extract_from_image(self, file_path: Path) -> str:
        """从图片提取文本（OCR）"""
        if not self.ocr_available:
            raise RuntimeError("OCR (PaddleOCR) is not available")

        try:
            result = self.paddle_ocr.ocr(str(file_path), cls=True)

            text_parts = []
            if result and result[0]:
                for line in result[0]:
                    if line and len(line) > 1:
                        text_parts.append(line[1][0])

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error performing OCR on {file_path}: {e}")
            raise

    def split_text_into_chunks(
        self,
        text: str,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ) -> List[str]:
        """
        将文本切分成块

        Args:
            text: 输入文本
            chunk_size: 块大小
            chunk_overlap: 重叠大小

        Returns:
            文本块列表
        """
        chunk_size = chunk_size or self.chunk_size
        chunk_overlap = chunk_overlap or self.chunk_overlap

        if not text or not text.strip():
            return []

        # 清理文本
        text = self._clean_text(text)

        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            # 如果当前段落加上后超过块大小
            if len(current_chunk) + len(paragraph) + 2 > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # 保留重叠部分
                if chunk_overlap > 0:
                    overlap_text = current_chunk[-chunk_overlap:] if len(current_chunk) > chunk_overlap else current_chunk
                    current_chunk = overlap_text + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph

        # 添加最后一个块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # 处理超长段落
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > chunk_size * 1.5:  # 如果块太长
                sub_chunks = self._split_long_chunk(chunk, chunk_size, chunk_overlap)
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)

        # 过滤掉太小的块
        return [
            chunk for chunk in final_chunks
            if len(chunk) >= self.min_chunk_size
        ]

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        # 压缩多个空格为一个
        text = re.sub(r' +', ' ', text)
        return text

    def _split_long_chunk(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """切分超长文本块"""
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            # 尝试在句子边界分割
            if end < len(text):
                # 查找最近的句号、问号、感叹号
                for punctuation in ['。', '！', '？', '.', '!', '?', '\n']:
                    last_punct = text.rfind(punctuation, start, end)
                    if last_punct > start:
                        end = last_punct + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - chunk_overlap if chunk_overlap > 0 else end

        return chunks

    def process_file(
        self,
        file_path: str,
        metadata: Dict[str, Any] = None,
    ) -> Tuple[List[str], List[Dict[str, Any]]]:
        """
        处理单个文件：提取文本并切分

        Args:
            file_path: 文件路径
            metadata: 基础元数据

        Returns:
            (文档块列表, 元数据列表)
        """
        if metadata is None:
            metadata = {}

        # 添加文件信息到元数据
        file_path_obj = Path(file_path)
        metadata["source_file"] = file_path_obj.name
        metadata["file_type"] = file_path_obj.suffix

        # 提取文本
        text = self.extract_text_from_file(file_path)

        # 切分文本
        chunks = self.split_text_into_chunks(text)

        # 为每个块创建元数据
        metadatas = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_index"] = i
            chunk_metadata["total_chunks"] = len(chunks)
            chunk_metadata["chunk_length"] = len(chunk)
            metadatas.append(chunk_metadata)

        logger.info(f"Processed {file_path}: {len(chunks)} chunks")
        return chunks, metadatas

    def process_directory(
        self,
        directory: str,
        extensions: List[str] = None,
        metadata: Dict[str, Any] = None,
    ) -> Tuple[List[str], List[Dict[str, Any]]]:
        """
        处理整个目录

        Args:
            directory: 目录路径
            extensions: 文件扩展名过滤（可选）
            metadata: 基础元数据

        Returns:
            (所有文档块列表, 所有元数据列表)
        """
        directory = Path(directory)
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory: {directory}")

        if extensions:
            extensions = set(ext.lower() for ext in extensions)

        all_chunks = []
        all_metadatas = []

        # 遍历目录
        for file_path in directory.rglob("*"):
            if file_path.is_file():
                # 检查扩展名
                if extensions and file_path.suffix.lower() not in extensions:
                    continue

                # 检查是否支持
                if not self.is_supported_format(str(file_path)):
                    continue

                try:
                    chunks, metadatas = self.process_file(str(file_path), metadata)
                    all_chunks.extend(chunks)
                    all_metadatas.extend(metadatas)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    continue

        logger.info(f"Processed directory {directory}: {len(all_chunks)} total chunks")
        return all_chunks, all_metadatas
